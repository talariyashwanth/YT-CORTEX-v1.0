"""Document loaders for supported file formats."""

from __future__ import annotations

import re
import uuid
from pathlib import Path

import fitz  # PyMuPDF

from src import config
from src.models.schemas import ExtractedDocument, PageContent


class IngestionError(Exception):
    """Raised when document ingestion fails."""


def validate_file(path: Path | str) -> Path:
    file_path = Path(path)
    if not file_path.exists():
        raise IngestionError(f"File not found: {file_path}")
    suffix = file_path.suffix.lower()
    if suffix not in config.SUPPORTED_EXTENSIONS:
        raise IngestionError(
            f"Unsupported format '{suffix}'. "
            f"Supported: {', '.join(sorted(config.SUPPORTED_EXTENSIONS))}"
        )
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > config.MAX_FILE_SIZE_MB:
        raise IngestionError(
            f"File exceeds {config.MAX_FILE_SIZE_MB} MB limit ({size_mb:.1f} MB)."
        )
    return file_path


def _detect_scanned_pdf(doc: fitz.Document) -> bool:
    """Heuristic: pages with very little extractable text may be scanned."""
    low_text_pages = 0
    for page in doc:
        if len(page.get_text().strip()) < 50:
            low_text_pages += 1
    return low_text_pages > len(doc) * 0.5 if len(doc) > 0 else False


def _extract_headings(text: str) -> str:
    lines = text.split("\n")
    for line in lines[:5]:
        stripped = line.strip()
        if stripped and (stripped.isupper() or re.match(r"^#{1,3}\s", stripped)):
            return stripped.lstrip("#").strip()
    return ""


def load_pdf(path: Path, document_id: str | None = None) -> ExtractedDocument:
    doc = fitz.open(path)
    pages: list[PageContent] = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            pages.append(
                PageContent(
                    page_number=i + 1,
                    text=text,
                    section_title=_extract_headings(text),
                )
            )
    is_scanned = _detect_scanned_pdf(doc)
    doc.close()

    if not pages:
        raise IngestionError("No extractable text found in PDF.")

    return ExtractedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        file_type=".pdf",
        pages=pages,
        metadata={"page_count": len(pages)},
        is_scanned_warning=is_scanned,
    )


def load_txt(path: Path, document_id: str | None = None) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        raise IngestionError("Text file is empty.")
    return ExtractedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        file_type=".txt",
        pages=[PageContent(page_number=1, text=text, section_title=path.stem)],
        metadata={"page_count": 1},
    )


def load_markdown(path: Path, document_id: str | None = None) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        raise IngestionError("Markdown file is empty.")
    sections = re.split(r"\n(?=#{1,3}\s)", text)
    pages: list[PageContent] = []
    for i, section in enumerate(sections):
        section = section.strip()
        if not section:
            continue
        title = _extract_headings(section)
        pages.append(PageContent(page_number=i + 1, text=section, section_title=title))
    if not pages:
        pages = [PageContent(page_number=1, text=text, section_title=path.stem)]
    return ExtractedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        file_type=path.suffix.lower(),
        pages=pages,
        metadata={"page_count": len(pages)},
    )


def load_docx(path: Path, document_id: str | None = None) -> ExtractedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise IngestionError("No extractable text found in DOCX.")

    full_text = "\n\n".join(paragraphs)
    section_title = ""
    for p in doc.paragraphs[:5]:
        if p.style and p.style.name and "Heading" in p.style.name:
            section_title = p.text.strip()
            break

    return ExtractedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        file_type=".docx",
        pages=[PageContent(page_number=1, text=full_text, section_title=section_title)],
        metadata={"page_count": 1, "paragraph_count": len(paragraphs)},
    )


def load_document(path: Path | str, document_id: str | None = None) -> ExtractedDocument:
    """Load and extract text from a supported document."""
    file_path = validate_file(path)
    suffix = file_path.suffix.lower()
    loaders = {
        ".pdf": load_pdf,
        ".txt": load_txt,
        ".md": load_markdown,
        ".markdown": load_markdown,
        ".docx": load_docx,
    }
    loader = loaders.get(suffix)
    if loader is None:
        raise IngestionError(f"No loader for format: {suffix}")
    return loader(file_path, document_id=document_id)
